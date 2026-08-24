"""Telegram front-end for the Career Agent (runs on OpenCode + OpenRouter)."""
import asyncio
import datetime as dt
import html
import json
import logging
import re
import secrets
import shutil
from pathlib import Path
from zoneinfo import ZoneInfo

from telegram import (InlineKeyboardButton, InlineKeyboardMarkup,
                      ReplyParameters, Update)
from telegram.error import NetworkError, TimedOut
from telegram.ext import (Application, CallbackQueryHandler, CommandHandler,
                          ContextTypes, MessageHandler, filters)

import classify
import config
import extract
import jobs_store
import onboarding
import render
import preferences
import scan
import telegram_format
import threads
from agent import run_turn

# Plain-text uploads OpenCode's read tool handles fine via --file.
NATIVE_READ_EXT = {".txt", ".md", ".markdown", ".csv"}
# Image uploads we transcribe server-side (extract.extract_image).
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp"}

# Only these get sent back to the user as generated documents.
SEND_BACK_EXT = {".md", ".pdf", ".docx", ".txt"}

_SECRET_PATTERNS = [
    re.compile(r"/bot\d{6,12}:[A-Za-z0-9_\-]{20,}"),             # token in a URL path
    re.compile(r"\bsk-[A-Za-z0-9_\-]{16,}"),                     # OpenAI/OpenRouter keys
    re.compile(r"\b\d{8,10}:[A-Za-z0-9_\-]{30,}"),               # Telegram bot token
    re.compile(r"\bBearer\s+[A-Za-z0-9._\-]{16,}", re.I),        # Authorization headers
    re.compile(r"(?i)\b([A-Z0-9_]*(?:KEY|TOKEN|SECRET|PASSWORD))\b\s*[=:]\s*\S+"),
]


def _scrub(text: str) -> str:
    """Replace anything that looks like a credential with <redacted>."""
    for pat in _SECRET_PATTERNS:
        if pat.pattern.startswith("/bot"):
            text = pat.sub("/bot<redacted>", text)
        else:
            text = pat.sub(
                lambda m: (m.group(1) + "=<redacted>") if m.groups() else "<redacted>",
                text)
    return text


class _RedactingFilter(logging.Filter):
    """Keep credentials out of the logs.

    httpx logs every request URL at INFO, and the Telegram bot token lives in
    the URL PATH ("/bot<token>/sendMessage"), so the token was printed on every
    single API call — a `docker logs` away from anyone with shell access. The
    plain token pattern does not catch it there: "bot" and the leading digit
    are both word characters, so there is no \\b for it to anchor on. Hence the
    dedicated /bot… pattern above.

    The URL arrives as a %s ARGUMENT, and not as a string either — httpx logs
    an httpx.URL object. A first cut scrubbed record.msg plus any str args and
    still leaked the token in production, because the URL is neither. So render
    the message the way the handler will and scrub THAT: it catches a secret
    wherever it hides and whatever type carries it.
    """

    def filter(self, record: logging.LogRecord) -> bool:
        try:
            message = record.getMessage()
        except Exception:  # noqa: BLE001 — a bad format string is not our problem
            return True
        scrubbed = _scrub(message)
        if scrubbed != message:
            # Args are folded into the rendered text, so they must be cleared
            # or the handler would try to interpolate a second time.
            record.msg, record.args = scrubbed, ()
        return True


logging.basicConfig(
    format="%(asctime)s %(levelname)s %(name)s: %(message)s", level=logging.INFO)
# On the handlers, not the root logger: a Filter on a logger is not consulted
# for records that propagate up from child loggers, so an httpx record would
# have sailed straight past it.
for _h in logging.getLogger().handlers:
    _h.addFilter(_RedactingFilter())
log = logging.getLogger("career-agent")

_STATIC_INTRO = (
    "👋 I'm your Career Agent.\n\n"
    "Send me:\n"
    "• your existing resume as a file (PDF / DOCX) or a photo → I read it and pull your real experience into memory\n"
    "• an experience or project as text → I structure it into a CV point and remember it\n"
    "• a job description link or pasted text → I assess your fit and draft a tailored resume\n"
    "• your goals, vision, or long-term plans → I store them\n\n"
    "I will never invent experience you don't have.\n\n"
    "Commands: /help  /reset  /onboard")


# --- Per-thread session id (OpenCode conversation continuity) --------------
def _session_path(chat_id: int, thread: str = threads.MAIN) -> Path:
    # The main thread keeps the pre-thread filename so an existing install's
    # conversation survives this upgrade without a migration step.
    if thread == threads.MAIN:
        return config.SESSIONS_DIR / f"{chat_id}.json"
    return config.SESSIONS_DIR / f"{chat_id}.{thread}.json"


def load_session_id(chat_id: int, thread: str = threads.MAIN):
    p = _session_path(chat_id, thread)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8")).get("session_id")
        except Exception:  # noqa: BLE001
            return None
    return None


def save_session_id(chat_id: int, session_id, thread: str = threads.MAIN) -> None:
    if session_id:
        _session_path(chat_id, thread).write_text(
            json.dumps({"session_id": session_id}), encoding="utf-8")


# One lock per (chat, thread). A turn is load-session -> long await -> save,
# so two concurrent turns on the SAME session interleave and the second write
# wins, silently dropping the first turn's history. Different threads hold
# different locks and stay independent, which is the whole point of threading.
_thread_locks: dict = {}


def _lock_for(chat_id: int, thread: str) -> asyncio.Lock:
    return _thread_locks.setdefault((chat_id, thread), asyncio.Lock())


def _context_stats(chat_id: int) -> dict:
    """Return whether a conversation is currently active.

    OpenCode stores session state in its own location, so we report only
    whether a session is active rather than guessing a byte count.
    """
    return {"active": bool(load_session_id(chat_id)), "sized": False}


def _status_text(chat_id: int) -> str:
    s = _context_stats(chat_id)
    if not s["active"]:
        return ("📊 <b>Conversation status</b>\n\n"
                "🟢 Fresh — no active conversation yet. The next message starts one.\n\n"
                "Your long-term memory (profile, goals, experiences) is separate and "
                "always kept.")
    return ("📊 <b>Conversation status</b>\n\n"
            "🟢 Active conversation.\n"
            "A long chat can still slow replies — reset anytime to start fresh.\n\n"
            "Your long-term memory (profile, goals, experiences) is <b>separate</b> "
            "and untouched by a reset.\n\n"
            "Tap below to clear the conversation and start fresh.")


def _status_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [[InlineKeyboardButton("🧹 Reset context", callback_data="reset_context")]])


# --- Critique-it button ----------------------------------------------------
_critique_tokens: dict[str, str] = {}  # token -> resume filename (in-memory)


def _register_critique(name: str) -> str:
    """Map a collision-resistant token to a resume filename for a 'Critique it'
    button. Tokens are random (not a restart-resettable counter) so a button
    from a previous process never resolves to a different resume after a
    restart — it misses the map and hits the graceful 'tap expired' path."""
    token = secrets.token_urlsafe(6)
    _critique_tokens[token] = name
    return token


def _critique_keyboard(token: str) -> InlineKeyboardMarkup:
    """Return a one-button keyboard for the Critique-it offer."""
    return InlineKeyboardMarkup(
        [[InlineKeyboardButton("📝 Critique it", callback_data=f"crit:{token}")]])


# --- Job-thread action buttons ---------------------------------------------
# Answering "yes" to "want me to build a resume?" used to mean typing it, which
# is the moment the reply-gesture rule bit hardest. Offer the two real choices
# as buttons under the answer instead.
_in_flight_threads: set = set()   # threads currently generating a resume


def _job_actions_keyboard(thread: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[
        InlineKeyboardButton("📄 Create resume", callback_data=f"th:cv:{thread}"),
        InlineKeyboardButton("⏭ Skip", callback_data=f"th:sk:{thread}"),
    ]])


def _job_actions_for(chat_id: int, thread: str):
    """Buttons to offer under a reply, or None.

    Only in a job thread, and only until that job has a resume — after that the
    Critique button on the delivered file is the useful next step, and repeating
    "Create resume" under every follow-up would just be noise.
    """
    if not thread or thread == threads.MAIN:
        return None
    meta = threads.get(chat_id, thread)
    if not meta or meta.get("resume"):
        return None
    return _job_actions_keyboard(thread)


def _clear_session(chat_id: int, thread: str = threads.MAIN) -> None:
    p = _session_path(chat_id, thread)
    if p.exists():
        p.unlink()


def _clear_all_sessions(chat_id: int) -> None:
    """Reset every thread for this chat, not just the main one.

    /reset means "start fresh"; leaving job threads pinned to old sessions
    would keep serving the very history the user asked to be rid of.
    """
    for key, _meta in threads.listing(chat_id):
        _clear_session(chat_id, key)
    _clear_session(chat_id, threads.MAIN)
    threads.forget_all(chat_id)


def _allowed(update: Update) -> bool:
    if not config.ALLOWED_USER_IDS:
        return True
    return bool(update.effective_user
                and update.effective_user.id in config.ALLOWED_USER_IDS)


# Anything that looks like a credential, redacted before an error goes out over
# Telegram. The chat is the operator's own diagnostic channel — surfacing the
# real failure is the whole point of raising instead of returning a placeholder
# — but an error string is attacker-influenced (it can carry provider response
# bodies and opencode stderr), so it must never be able to smuggle a key out.
_MAX_ERROR_CHARS = 400


def _safe_error(e) -> str:
    """Render an exception for the user: credentials scrubbed, length bounded."""
    text = _scrub(str(e).strip() or e.__class__.__name__)
    if len(text) > _MAX_ERROR_CHARS:
        text = text[:_MAX_ERROR_CHARS] + "… (truncated — full detail in the logs)"
    return text


def _message_id(sent):
    """Message id of a python-telegram-bot send result, or None.

    Test doubles return plain objects (or nothing), so this never assumes the
    real Message type.
    """
    return getattr(sent, "message_id", None)


def _thread_prefix(chat_id: int, thread: str) -> str:
    """A '🧵 label' header so an answer is visibly tied to its job.

    A Telegram DM is one flat stream with no real threads, so the per-job
    sessions are otherwise invisible: two jobs in flight read as one jumbled
    conversation even though the model keeps them apart. Markdown-active
    characters are stripped from the label — it goes through the Markdown ->
    HTML converter, and a company name containing '*' would eat the formatting.
    """
    if not thread or thread == threads.MAIN:
        return ""
    meta = threads.get(chat_id, thread) or {}
    label = re.sub(r"[*_`\[\]]", "", meta.get("label") or thread).strip()
    return f"🧵 **{label}**\n\n" if label else ""


def _retry_prefix(chat_id: int, thread: str):
    """Context to re-prime a fresh session with when a thread's history is lost.

    OpenCode's stale-session retry starts from nothing, so a job thread would
    come back asking for a JD the user already sent. The label and source URL
    survive in the thread store, which is enough for the agent to re-read the
    posting itself.
    """
    if not thread or thread == threads.MAIN:
        return None
    meta = threads.get(chat_id, thread) or {}
    label, url = meta.get("label"), meta.get("url")
    if not label and not url:
        return None
    where = f" The posting is at {url} — re-read it if you need the details." if url else ""
    return (f"[Context: this conversation is about the job '{label}'. Its earlier "
            f"history was lost to an error, so you are starting fresh.{where}]")


def _session_was_reset(previous, current) -> bool:
    """True when a turn abandoned the session it was asked to resume."""
    return bool(previous and current and previous != current)


def _lost_history_notice(chat_id: int, thread: str) -> str:
    meta = threads.get(chat_id, thread) or {}
    label = meta.get("label") or "this job"
    return (f"⚠️ _Heads up: the previous run for **{label}** failed and I lost this "
            f"thread's history, so I'm working from the posting rather than our "
            f"earlier conversation._\n\n")


def _reply_params(reply_to):
    """Quote the message being answered, degrading if it has been deleted.

    allow_sending_without_reply keeps a deleted target from turning the whole
    reply into a BadRequest.
    """
    if not reply_to:
        return None
    return ReplyParameters(message_id=reply_to, allow_sending_without_reply=True)


async def _send_chat(bot, chat_id: int, text: str, thread: str = None,
                     reply_to: int = None, reply_markup=None) -> list:
    """Send text as Telegram HTML (plain-text fallback); return sent message ids.

    When `thread` is given the reply carries its label and each sent message is
    bound to it, so a reply to any of them routes the next turn back into the
    same conversation. `reply_to` quotes the message being answered — only on
    the first chunk, so a long answer isn't a wall of repeated quote bars.
    `reply_markup` rides on the LAST chunk, so buttons sit under the end of the
    answer rather than in the middle of it.
    """
    text = _thread_prefix(chat_id, thread) + (text or "")
    pieces = list(telegram_format.chunk(text)) or [""]
    ids = []
    params = _reply_params(reply_to)
    for i, piece in enumerate(pieces):
        kw = {}
        if params:
            kw["reply_parameters"] = params
        if reply_markup is not None and i == len(pieces) - 1:
            kw["reply_markup"] = reply_markup
        try:
            sent = await bot.send_message(
                chat_id, telegram_format.to_telegram_html(piece),
                parse_mode="HTML", **kw)
        except Exception as e:  # noqa: BLE001 - bad markup etc.: degrade gracefully
            log.warning("HTML send failed (%s); falling back to plain text", e)
            sent = await bot.send_message(chat_id, telegram_format.to_plain(piece), **kw)
        params = None  # subsequent chunks follow on naturally
        mid = _message_id(sent)
        if mid:
            ids.append(mid)
    if thread and thread != threads.MAIN:
        for mid in ids:
            threads.bind_message(chat_id, mid, thread)
    return ids


async def _send(update: Update, text: str, thread: str = None,
                reply_to: int = None) -> list:
    return await _send_chat(update.get_bot(), update.effective_chat.id, text,
                            thread, reply_to)


def _resume_snapshot() -> dict:
    """Map filename -> mtime, so we catch both new AND updated files."""
    snap = {}
    for p in config.RESUMES_DIR.glob("*"):
        try:
            snap[p.name] = p.stat().st_mtime
        except OSError:
            pass
    return snap


# --- Handlers --------------------------------------------------------------
async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    uid = update.effective_user.id if update.effective_user else "unknown"
    await update.message.reply_text(
        f"Your Telegram user ID is: {uid}\n"
        "(Put this in ALLOWED_USER_IDS in .env to keep the bot private.)")
    if onboarding.is_fresh() and onboarding.status() not in ("in_progress", "done"):
        await _launch_onboarding(update, ctx)
        return
    await update.message.reply_text(_STATIC_INTRO)


async def help_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "How to use me:\n\n"
        "📎 Send your resume: upload a PDF/DOCX file or a photo of it — I read it and store your real experience.\n"
        "📌 Build your memory: tell me about your role, goals, and what you've worked on.\n"
        "🧱 Add experience: describe something you did — I'll structure it (Situation/Task/Action/Result/Metrics) and save it.\n"
        "🎯 Check a job: paste a JD or a link. I'll rate the fit and name the gaps.\n"
        "📄 Tailored resume: ask for a resume for a JD — built only from what I actually know about you.\n\n"
        "🧵 Send several job links at once — each opens its own thread, so they never "
        "mix. Reply to a message to continue that job; /threads lists them.\n\n"
        "/status shows how big the current conversation is, with a one-tap Reset button.\n"
        "/reset starts a fresh conversation (your saved profile, goals, experiences and projects are kept).")


async def reset(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    _clear_all_sessions(update.effective_chat.id)
    await update.message.reply_text(
        "🧹 Started a fresh conversation and closed every job thread. "
        "Your long-term memory is untouched.")


async def onboard_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    await _launch_onboarding(update, ctx)


async def status(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    await update.message.reply_text(
        _status_text(update.effective_chat.id),
        parse_mode="HTML", reply_markup=_status_keyboard())


async def on_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    if not _allowed(update):
        await query.answer("Private bot.", show_alert=True)
        return
    data = query.data or ""
    if data.startswith("job:"):
        await _on_job_action(update, ctx, data)
        return
    if data.startswith("crit:"):
        await _on_critique_action(update, ctx, data)
        return
    if data.startswith("th:"):
        await _on_thread_action(update, ctx, data)
        return
    if query.data == "reset_context":
        _clear_all_sessions(update.effective_chat.id)
        await query.answer("Conversation cleared ✅")
        try:
            await query.edit_message_text(
                "🧹 <b>Conversation cleared.</b> Started fresh — your long-term "
                "memory (profile, goals, experiences) is untouched.",
                parse_mode="HTML")
        except Exception:  # noqa: BLE001 - message may be unchanged/too old to edit
            await query.message.reply_text("🧹 Conversation cleared. Memory kept.")


async def _keep_typing(bot, chat_id: int) -> None:
    """Re-send the 'typing…' action every few seconds during a long turn."""
    try:
        while True:
            await bot.send_chat_action(chat_id=chat_id, action="typing")
            await asyncio.sleep(4)
    except asyncio.CancelledError:
        pass


async def _launch_onboarding(update: Update, ctx: ContextTypes.DEFAULT_TYPE,
                             user_message: str = None) -> None:
    """Set status in_progress and run the kickoff turn in the chat's session.

    On agent failure, reset to not_started and fall back to the static intro so
    the user is never stuck mid-launch.
    """
    chat_id = update.effective_chat.id
    onboarding.set_status("in_progress")
    session_id = load_session_id(chat_id)
    typing = asyncio.create_task(_keep_typing(ctx.bot, chat_id))
    try:
        prompt = onboarding.ONBOARDING_KICKOFF
        if user_message:
            prompt += (f"\n\nThe user's first message to you was: {user_message!r} — "
                       "acknowledge it and weave it into the onboarding naturally.")
        text, session_id = await run_turn(prompt, session_id)
    except Exception:  # noqa: BLE001
        log.exception("onboarding kickoff failed")
        onboarding.set_status("not_started")
        await _send(update, _STATIC_INTRO)
        return
    finally:
        typing.cancel()
    save_session_id(chat_id, session_id)
    text, completed = onboarding.strip_complete_marker(text)
    if completed:
        onboarding.set_status("done")
    await _send(update, text)


async def _model_for_message(text: str) -> str | None:
    """Pick the model for a TYPED message. None = default model.

    Only classifies when routing is active (a distinct critique/resume model is
    configured), so there's no classifier cost unless the user opted in.
    """
    if not config.routing_active():
        return None
    # Onboarding turns run on the default model (spec: onboarding stays on default).
    if onboarding.status() == "in_progress":
        return None
    return config.model_for(await classify.classify_task(text))


async def _run_and_reply(update: Update, ctx: ContextTypes.DEFAULT_TYPE,
                         prompt: str, files=None, thread: str = None) -> None:
    """Run one agent turn in `thread` and send back text + any new resume.

    `files` are upload paths passed to OpenCode via --file. `thread` defaults
    to the main conversation; a job thread keeps its own OpenCode session so
    two jobs in flight never share a history.
    """
    chat_id = update.effective_chat.id
    thread = thread or threads.MAIN
    model = await _model_for_message(prompt)

    # Bind the user's own message too, so replying to it (not just to the
    # bot's answer) also routes back into this thread.
    if thread != threads.MAIN and update.message:
        threads.bind_message(chat_id, update.message.message_id, thread)

    async with _lock_for(chat_id, thread):
        previous = load_session_id(chat_id, thread)
        before = _resume_snapshot()
        typing = asyncio.create_task(_keep_typing(ctx.bot, chat_id))
        try:
            text, session_id = await run_turn(
                prompt, previous, files=files, model=model,
                retry_prefix=_retry_prefix(chat_id, thread))
        except Exception as e:  # noqa: BLE001
            log.exception("turn failed")
            sent = await update.message.reply_text(f"⚠️ Something went wrong: {_safe_error(e)}")
            if thread != threads.MAIN:
                threads.bind_message(chat_id, _message_id(sent), thread)
            return
        finally:
            typing.cancel()

        save_session_id(chat_id, session_id, thread)
        threads.touch(chat_id, thread)
        lost = _session_was_reset(previous, session_id)

    text, completed = onboarding.strip_complete_marker(text)
    if completed and onboarding.status() == "in_progress":
        onboarding.set_status("done")
    text, claimed = strip_resume_marker(text)
    text, job_label = strip_job_marker(text)
    if job_label and thread != threads.MAIN:
        threads.set_label(chat_id, thread, job_label)
    if lost:
        text = _lost_history_notice(chat_id, thread) + text
    # Quote the message being answered, so the job an answer belongs to is
    # visible in a chat that has no real threads.
    anchor = update.message.message_id if update.message else None
    await _send_chat(ctx.bot, chat_id, text, thread, reply_to=anchor,
                     reply_markup=_job_actions_for(chat_id, thread))
    await _deliver_new_files(update, before, thread, claimed, reply_to=anchor)


async def _send_doc_chat(bot, chat_id: int, path: Path, reply_markup=None,
                         reply_to: int = None) -> None:
    try:
        params = _reply_params(reply_to)
        kw = {"reply_parameters": params} if params else {}
        with open(path, "rb") as fh:
            await bot.send_document(
                chat_id, document=fh, filename=path.name,
                reply_markup=reply_markup, **kw)
    except Exception as e:  # noqa: BLE001
        log.warning("could not send %s: %s", path, e)


async def _send_doc(update: Update, path: Path) -> None:
    await _send_doc_chat(update.get_bot(), update.effective_chat.id, path)


JOB_MARKER_RE = re.compile(r"\[\[JOB:\s*([^\]]*?)\s*\]\]")

# Appended to any prompt that may identify a job. A link gives us only the
# hostname, so without this a thread stays labelled "linkedin.com" — useless
# once two of them are open.
JOB_MARKER_INSTRUCTION = (
    "If this message is about a specific job opening, identify it on the last "
    "line of your reply with this marker, using the job's real title and "
    "employer from the posting (never guess one): [[JOB:position|company]]")


def strip_job_marker(text: str):
    """Pull the agent's `[[JOB:position|company]]` claim out of a reply.

    Returns (clean_text, label or None). The pipe is an explicit separator
    rather than a dash the model would render inconsistently (-, --, —).
    """
    m = JOB_MARKER_RE.search(text or "")
    if not m:
        return (text or "").strip(), None
    raw = m.group(1).strip()
    clean = JOB_MARKER_RE.sub("", text).strip()
    if not raw:
        return clean, None
    position, _, company = raw.partition("|")
    return clean, (threads.format_label(position, company) or None)


RESUME_MARKER_RE = re.compile(r"\[\[RESUME:\s*([^\]\s][^\]]*?)\s*\]\]")

# Appended to any prompt that may produce a resume. The agent naming the file
# is what lets delivery send exactly that one, instead of guessing from a
# directory diff that cannot tell two concurrent jobs apart.
RESUME_MARKER_INSTRUCTION = (
    "When you have saved the resume JSON, end your reply with this marker on its "
    "own last line, using the exact filename you wrote (no path): "
    "[[RESUME:filename.json]]")


def strip_resume_marker(text: str):
    """Pull the agent's `[[RESUME:name.json]]` claim out of a reply.

    Returns (clean_text, filename or None). Without this, delivery can only
    guess which file a turn produced by diffing the whole resumes/ directory —
    which silently cross-delivers as soon as two threads run at once. Mirrors
    onboarding.strip_complete_marker.
    """
    m = RESUME_MARKER_RE.search(text or "")
    if not m:
        return (text or "").strip(), None
    name = Path(m.group(1).strip()).name   # never let the agent escape resumes/
    clean = RESUME_MARKER_RE.sub("", text).strip()
    return clean, (name or None)


def _deliverable_here(chat_id: int, name: str, thread: str, claimed: str) -> bool:
    """True if `name` belongs in this thread's chat.

    A file another thread already owns is never re-sent here — that is the
    cross-delivery this whole mechanism exists to prevent. Unowned files still
    go to whoever is running, so a turn that forgets the marker degrades to the
    old behaviour instead of delivering nothing.
    """
    if claimed and name == claimed:
        return True
    owner = threads.resume_owner(chat_id, name)
    return owner is None or owner == (thread or threads.MAIN)


async def _deliver_changed_resumes(bot, chat_id: int, before: dict,
                                   thread: str = None, claimed: str = None,
                                   reply_to: int = None) -> None:
    """Render any new/updated resume JSON to PDF and send it; send other docs.

    Compares mtimes so an *updated* resume (same filename) is re-rendered and
    re-sent. Skips scratch/helper files (e.g. _make_pdf.py). `claimed` is the
    filename the agent named via [[RESUME:...]]; it is delivered even if its
    mtime is unchanged, and it claims ownership for this thread.
    """
    after = _resume_snapshot()
    changed = sorted(name for name, mtime in after.items()
                     if not name.startswith((".", "_"))
                     and mtime != before.get(name))
    # A rewrite landing inside the same mtime tick still needs delivering.
    if claimed and claimed in after and claimed not in changed:
        changed.append(claimed)
    if claimed and threads.resume_owner(chat_id, claimed) is None:
        threads.set_resume(chat_id, thread or threads.MAIN, claimed)
    skipped = [n for n in changed if not _deliverable_here(chat_id, n, thread, claimed)]
    if skipped:
        log.info("thread %s: not re-sending %s (owned by another thread)",
                 thread, ", ".join(skipped))
    changed = [n for n in changed if _deliverable_here(chat_id, n, thread, claimed)]
    for name in changed:
        path = config.RESUMES_DIR / name
        ext = path.suffix.lower()
        if ext == ".json":
            token = _register_critique(name)
            kb = _critique_keyboard(token)
            button_sent = False
            try:
                pdf = await asyncio.to_thread(render.render_json_to_pdf, path)
                await _send_doc_chat(bot, chat_id, pdf, reply_markup=kb,
                                     reply_to=reply_to)
                # _send_doc_chat swallows its own send errors, so button_sent means
                # "render succeeded and the send call didn't raise" — not a delivery guarantee.
                button_sent = True
            except Exception as e:  # noqa: BLE001
                log.warning("PDF render failed for %s: %s", name, e)
                await bot.send_message(
                    chat_id,
                    "⚠️ I built your resume but couldn't render the PDF. "
                    "Sending the data file instead.")
            # JSON Resume file (portable). If the PDF didn't go out, the button
            # rides on the JSON so the one-tap critique is never lost.
            await _send_doc_chat(
                bot, chat_id, path, reply_markup=None if button_sent else kb,
                reply_to=reply_to)
        elif ext in SEND_BACK_EXT:
            await _send_doc_chat(bot, chat_id, path, reply_to=reply_to)


async def _deliver_new_files(update: Update, before: dict, thread: str = None,
                             claimed: str = None, reply_to: int = None) -> None:
    await _deliver_changed_resumes(
        update.get_bot(), update.effective_chat.id, before, thread, claimed,
        reply_to)


# --- Job discovery ---------------------------------------------------------
try:
    _SCAN_TZ = ZoneInfo(config.SCAN_TZ)  # reused by scheduler + scan-day gate
except Exception:  # noqa: BLE001 - bad SCAN_TZ shouldn't crash the whole bot
    log.warning("Invalid SCAN_TZ %r — falling back to UTC", config.SCAN_TZ)
    _SCAN_TZ = ZoneInfo("UTC")
_in_flight_applies: set = set()  # job ids currently generating a resume (double-tap guard)
_pending_skip_reason: dict = {}  # chat_id -> jid awaiting a free-text skip reason
_FALLBACK_SKIP_REASONS = ["Too senior", "Too junior", "Location", "Wrong tech"]  # capped at 4 (one row)


def _skip_reason_keyboard(jid: str, job: dict) -> InlineKeyboardMarkup:
    reasons = (job.get("skip_reasons") or _FALLBACK_SKIP_REASONS)[:4]
    rows, row = [], []
    for i, r in enumerate(reasons):
        row.append(InlineKeyboardButton(r[:24], callback_data=f"job:sk:{jid}:{i}"))
        if len(row) == 2:
            rows.append(row)
            row = []
    if row:
        rows.append(row)
    rows.append([
        InlineKeyboardButton("✏️ Other…", callback_data=f"job:sk:{jid}:other"),
        InlineKeyboardButton("⏭ Skip", callback_data=f"job:sk:{jid}:none"),
    ])
    return InlineKeyboardMarkup(rows)


def _job_card_html(job: dict) -> str:
    title = html.escape(job.get("title", "Role"))
    company = html.escape(job.get("company", ""))
    location = html.escape(job.get("location", "") or "—")
    score = job.get("fit_score")
    score_line = f"Fit <b>{score}/10</b>" if score is not None else "Fit: n/a"
    why_fit = html.escape(job.get("why_fit", "") or "")
    why_aligns = html.escape(job.get("why_aligns", "") or "")
    url = html.escape(job.get("url", "") or "")
    if url and not url.lower().startswith(("https://", "http://")):
        url = ""  # drop non-http schemes (e.g. javascript:) before putting in href
    lines = [f"<b>{title}</b> @ {company} · {location}", score_line]
    if why_fit:
        lines.append(f"💪 {why_fit}")
    if why_aligns:
        lines.append(f"🎯 {why_aligns}")
    if url:
        lines.append(f'🔗 <a href="{url}">View posting</a>')
    return "\n".join(lines)


async def _send_job_card(bot, chat_id: int, job: dict) -> None:
    kb = InlineKeyboardMarkup([[
        InlineKeyboardButton("📄 Apply", callback_data=f"job:apply:{job['id']}"),
        InlineKeyboardButton("⏭ Skip", callback_data=f"job:skip:{job['id']}"),
    ]])
    await bot.send_message(
        chat_id, _job_card_html(job), parse_mode="HTML", reply_markup=kb)


async def _on_job_action(update: Update, ctx: ContextTypes.DEFAULT_TYPE, data: str) -> None:
    query = update.callback_query
    parts = data.split(":")
    if len(parts) < 3:
        await query.answer()
        return
    action, jid = parts[1], parts[2]
    extra = parts[3] if len(parts) > 3 else None
    _pending_skip_reason.pop(update.effective_chat.id, None)  # user re-engaged a card; drop any half-finished free-text reason
    job = jobs_store.get(jid)
    if not job:
        await query.answer("That job is no longer available.", show_alert=True)
        return

    if action == "skip":
        try:
            await query.edit_message_reply_markup(reply_markup=_skip_reason_keyboard(jid, job))
        except Exception:  # noqa: BLE001 - message too old / unchanged
            pass
        await query.answer("Why skip it?")
        return

    if action == "sk":
        if extra == "other":
            _pending_skip_reason[update.effective_chat.id] = jid
            await query.answer("Type your reason 👇")
            try:
                await query.edit_message_reply_markup(reply_markup=None)
            except Exception:  # noqa: BLE001
                pass
            await ctx.bot.send_message(
                update.effective_chat.id,
                "Reply with your reason for skipping (one line) and I'll learn from it.")
            return
        reason = None
        if extra is not None and extra.isdigit():
            reasons = job.get("skip_reasons") or _FALLBACK_SKIP_REASONS
            i = int(extra)
            if 0 <= i < len(reasons):
                reason = reasons[i]
        jobs_store.set_decision(jid, "skipped", reason)
        await query.answer("Skipped ⏭")
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001
            pass
        return

    if action == "apply":
        if jid in _in_flight_applies:
            await query.answer("Already generating — please wait.", show_alert=True)
            return
        await query.answer("Tailoring your resume…")
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001
            pass
        _in_flight_applies.add(jid)
        try:
            await _generate_resume_for(ctx, update.effective_chat.id, job, jid)
        finally:
            _in_flight_applies.discard(jid)
        return

    await query.answer()  # unknown action — dismiss the spinner


async def _generate_resume_for(ctx: ContextTypes.DEFAULT_TYPE, chat_id: int, job: dict, jid: str) -> None:
    # Its own thread: tapping Apply on three scan results gives three separate
    # conversations, so "make it shorter" is never ambiguous.
    label = threads.format_label(job.get("title"), job.get("company"))
    thread = threads.new_thread(chat_id, label or f"job {jid[:8]}",
                                named=bool(label), url=job.get("url"))
    prompt = (
        "The user chose to apply to this job from a discovery scan. Build a tailored "
        "resume for it, following ALL your resume rules (never fabricate).\n\n"
        f"Job title: {job.get('title')}\n"
        f"Company: {job.get('company')}\n"
        f"Location: {job.get('location')}\n"
        f"Link: {job.get('url')}\n\n"
        "If the link is reachable, WebFetch it to read the full JD; if it's blocked, "
        "tailor from the details above and the user's memory. Save the resume JSON to "
        "resumes/ as usual so the PDF is generated, then briefly tell me what you "
        "emphasized and any real gaps.\n\n" + RESUME_MARKER_INSTRUCTION
    )
    sent = await ctx.bot.send_message(
        chat_id,
        f"📄 Tailoring your resume for {job.get('title')} @ {job.get('company')} — about a minute…")
    threads.bind_message(chat_id, _message_id(sent), thread)

    async with _lock_for(chat_id, thread):
        session_id = load_session_id(chat_id, thread)
        before = _resume_snapshot()
        typing = asyncio.create_task(_keep_typing(ctx.bot, chat_id))
        try:
            text, session_id = await run_turn(prompt, session_id,
                                              model=config.model_for("resume"))
        except Exception as e:  # noqa: BLE001
            log.exception("resume generation failed")
            threads.forget(chat_id, thread)  # nothing happened in it; don't leave a stub
            await ctx.bot.send_message(
                chat_id,
                f"⚠️ Couldn't build the resume: {_safe_error(e)}\n\n"
                f"You can still apply — send me the job link and I'll tailor one.")
            return
        finally:
            typing.cancel()

        save_session_id(chat_id, session_id, thread)
        threads.touch(chat_id, thread)

    jobs_store.set_decision(jid, "applied")
    text, claimed = strip_resume_marker(text)
    text, _job_label = strip_job_marker(text)   # already named from the job card
    await _send_chat(ctx.bot, chat_id, text, thread,
                     reply_to=_message_id(sent))   # quote the "Tailoring…" notice
    await _deliver_changed_resumes(ctx.bot, chat_id, before, thread, claimed,
                                   reply_to=_message_id(sent))


async def _on_thread_action(update: Update, ctx: ContextTypes.DEFAULT_TYPE,
                            data: str) -> None:
    """Handle the Create-resume / Skip buttons under a job thread's reply."""
    query = update.callback_query
    parts = data.split(":")
    if len(parts) < 3:
        await query.answer()
        return
    action, thread = parts[1], parts[2]
    chat_id = update.effective_chat.id

    if not threads.exists(chat_id, thread) or thread == threads.MAIN:
        await query.answer("That job thread is closed.", show_alert=True)
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001 - message too old / unchanged
            pass
        return

    if action == "sk":
        label = (threads.get(chat_id, thread) or {}).get("label") or "that job"
        threads.forget(chat_id, thread)   # also drops it as the current thread
        _clear_session(chat_id, thread)
        await query.answer("Skipped ⏭")
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001
            pass
        await ctx.bot.send_message(
            chat_id, f"⏭ Closed <b>{html.escape(label)}</b>. Back in the main "
                     "conversation.", parse_mode="HTML")
        return

    if action == "cv":
        if thread in _in_flight_threads:
            await query.answer("Already building — please wait.", show_alert=True)
            return
        await query.answer("Building your resume…")
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001
            pass
        _in_flight_threads.add(thread)
        try:
            await _generate_resume_in_thread(
                ctx, chat_id, thread, _message_id(getattr(query, "message", None)))
        finally:
            _in_flight_threads.discard(thread)
        return

    await query.answer()  # unknown action — dismiss the spinner


async def _generate_resume_in_thread(ctx: ContextTypes.DEFAULT_TYPE, chat_id: int,
                                     thread: str, anchor: int = None) -> None:
    """Build a resume for the job this thread is already about.

    The thread's session holds the JD, so the prompt does not restate it —
    restating a half-remembered version is how invented details get in.
    """
    meta = threads.get(chat_id, thread) or {}
    label = meta.get("label") or "this job"
    prompt = (
        "Build a tailored resume for the job in this conversation, following ALL "
        "your resume rules (never fabricate; use only what memory actually holds). "
        "Save the resume JSON to resumes/ as usual so the PDF is generated, then "
        "briefly tell me what you emphasized and any real gaps.\n\n"
        + RESUME_MARKER_INSTRUCTION)
    sent = await ctx.bot.send_message(
        chat_id, f"📄 Building your resume for {label} — about a minute…")
    threads.bind_message(chat_id, _message_id(sent), thread)

    async with _lock_for(chat_id, thread):
        previous = load_session_id(chat_id, thread)
        before = _resume_snapshot()
        typing = asyncio.create_task(_keep_typing(ctx.bot, chat_id))
        try:
            text, session_id = await run_turn(
                prompt, previous, model=config.model_for("resume"),
                retry_prefix=_retry_prefix(chat_id, thread))
        except Exception as e:  # noqa: BLE001
            log.exception("in-thread resume generation failed")
            await ctx.bot.send_message(chat_id, f"⚠️ Couldn't build the resume: {_safe_error(e)}")
            return
        finally:
            typing.cancel()
        save_session_id(chat_id, session_id, thread)
        threads.touch(chat_id, thread)
        lost = _session_was_reset(previous, session_id)

    text, claimed = strip_resume_marker(text)
    text, _job = strip_job_marker(text)
    if lost:
        text = _lost_history_notice(chat_id, thread) + text
    await _send_chat(ctx.bot, chat_id, text, thread, reply_to=anchor)
    await _deliver_changed_resumes(ctx.bot, chat_id, before, thread, claimed,
                                   reply_to=anchor)


async def _on_critique_action(update: Update, ctx: ContextTypes.DEFAULT_TYPE,
                              data: str) -> None:
    """One-tap critique: score the just-generated resume on the critique model."""
    query = update.callback_query
    token = data.split(":", 1)[1] if ":" in data else ""
    name = _critique_tokens.get(token)
    if not name:  # map cleared by a restart, unknown/expired token, etc.
        await query.answer("Tap expired — just type 'critique it'.", show_alert=True)
        try:
            await query.edit_message_reply_markup(reply_markup=None)
        except Exception:  # noqa: BLE001 - message too old to edit
            pass
        return
    _critique_tokens.pop(token, None)  # single-use: bound the map + close the double-tap window
    # Consume the button first: this is also the double-tap guard.
    try:
        await query.edit_message_reply_markup(reply_markup=None)
    except Exception:  # noqa: BLE001
        pass
    await query.answer("Scoring your resume…")

    chat_id = update.effective_chat.id
    # Critique inside the thread that built this resume — that session is the
    # only one holding the JD the prompt refers to. Falls back to the main
    # conversation for resumes written before threads existed.
    thread = threads.resume_owner(chat_id, name) or threads.MAIN
    prompt = (
        f"Critique resumes/{name} and score it against the JD from this "
        "conversation, following your normal critique rules (the compact Telegram "
        "scorecard). If you can't tell which JD this resume targets, ask me to "
        "paste it rather than guessing."
    )
    typing = asyncio.create_task(_keep_typing(ctx.bot, chat_id))
    try:
        async with _lock_for(chat_id, thread):
            session_id = load_session_id(chat_id, thread)
            text, session_id = await run_turn(
                prompt, session_id, model=config.model_for("critique"))
            save_session_id(chat_id, session_id, thread)
    except Exception as e:  # noqa: BLE001
        log.exception("critique failed")
        await ctx.bot.send_message(chat_id, f"⚠️ Couldn't run the critique: {_safe_error(e)}")
        return
    finally:
        typing.cancel()
    # Quote the resume the button was attached to, so the score sits with it.
    await _send_chat(ctx.bot, chat_id, text, thread,
                     reply_to=_message_id(getattr(query, "message", None)))


async def _do_scan(bot, chat_id: int, manual: bool) -> None:
    try:
        note = await preferences.run_synthesis()
        if note:
            await bot.send_message(chat_id, f"🧠 {note}")
    except Exception:  # noqa: BLE001 - learning must never block a scan
        log.warning("preference synthesis failed", exc_info=True)

    try:
        matches = await scan.run_scan()
    except scan.ScanError as e:
        log.warning("scan failed: %s", e)
        if manual:
            await bot.send_message(chat_id, f"⚠️ Scan failed — {_safe_error(e)}. Try again later.")
        return
    if not matches:
        if manual or not config.SILENT_WHEN_EMPTY:
            await bot.send_message(chat_id, "🔍 No new strong matches this time.")
        return
    await bot.send_message(chat_id, f"🔍 Found {len(matches)} new match(es):")
    for job in matches:
        await _send_job_card(bot, chat_id, job)


def _owner_chat_id() -> int:
    if config.OWNER_CHAT_ID:
        return config.OWNER_CHAT_ID
    if len(config.ALLOWED_USER_IDS) == 1:
        return next(iter(config.ALLOWED_USER_IDS))
    return 0


async def _scheduled_scan(ctx: ContextTypes.DEFAULT_TYPE) -> None:
    """Daily timer; only actually scans on configured weekdays (Monday=0)."""
    today = dt.datetime.now(_SCAN_TZ).weekday()
    if today not in config.SCAN_WEEKDAYS:
        return
    chat_id = _owner_chat_id()
    if not chat_id:
        log.warning("Scheduled scan skipped: set OWNER_CHAT_ID in .env.")
        return
    await _do_scan(ctx.bot, chat_id, manual=False)


async def scan_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    await update.message.reply_text("🔍 Scanning for jobs… this can take a minute.")
    try:
        await _do_scan(ctx.bot, update.effective_chat.id, manual=True)
    except Exception as e:  # noqa: BLE001
        log.exception("scan_cmd failed")
        await update.message.reply_text(f"⚠️ Scan error — {_safe_error(e)}. Try again later.")


def _safe_name(name: str) -> str:
    base = Path(name or "").name or "upload"
    return re.sub(r"[^A-Za-z0-9._-]", "_", base)[:80]


def _docx_to_text(path: Path):
    try:
        import docx  # python-docx
    except ImportError:
        return None
    try:
        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip() or None
    except Exception:  # noqa: BLE001
        return None


def _route_thread(update: Update, allow_new: bool = True):
    """Decide which conversation a message belongs to.

    Returns (thread_key, opened) where `opened` is True for a brand-new thread.
    `allow_new=False` routes replies but never forks: an uploaded CV is a memory
    operation and belongs in the main conversation, unless it is a reply into a
    job thread (e.g. sending the JD as a PDF).

    Replying to a message keeps you in the thread that message came from — the
    native Telegram gesture, so no command to remember and the client shows the
    thread visually. A fresh (non-reply) message carrying a link starts its own
    thread, which is the case that used to turn one chat into a tangle: send
    three job links and each gets a clean conversation instead of all three
    sharing a history. Everything else stays in the main conversation.
    """
    chat_id = update.effective_chat.id
    msg = update.message
    # A document/photo message carries no .text at all — its words live in
    # .caption — so neither attribute can be assumed present.
    text = (getattr(msg, "text", None) or getattr(msg, "caption", None) or "") if msg else ""

    reply_to = getattr(msg, "reply_to_message", None) if msg else None
    if reply_to is not None:
        existing = threads.thread_for_message(chat_id, reply_to.message_id)
        if existing and threads.exists(chat_id, existing):
            threads.set_current(chat_id, existing)   # replying switches threads
            return existing, False

    if onboarding.status() == "in_progress":
        return threads.MAIN, False   # onboarding is one linear interview

    url = threads.find_url(text) if allow_new else None
    if url:
        return threads.new_thread(chat_id, threads.label_for_url(url), url=url), True

    # Sticky: keep typing in whichever job was last worked on. Requiring a reply
    # here is what sent a bare "yes" — answering the bot's own question — into
    # the main conversation, where it had no idea what had been agreed.
    return threads.current(chat_id), False


async def _announce_thread(bot, chat_id: int, thread: str) -> None:
    """Tell the user a separate conversation just opened, and how to stay in it."""
    meta = threads.get(chat_id, thread) or {}
    sent = await bot.send_message(
        chat_id,
        f"🧵 New thread for <b>{html.escape(meta.get('label') or thread)}</b> — "
        "reply to my messages here to keep this job separate from the rest. "
        "/threads lists them.",
        parse_mode="HTML")
    threads.bind_message(chat_id, _message_id(sent), thread)


async def threads_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    chat_id = update.effective_chat.id
    items = threads.listing(chat_id)
    if not items:
        await update.message.reply_text(
            "No job threads yet. Send me a job link and I'll open one — each link "
            "gets its own conversation so they never mix.")
        return
    cur = threads.current(chat_id)
    lines = ["🧵 <b>Your job threads</b>", ""]
    for key, meta in items:
        resume = meta.get("resume")
        tail = f" — 📄 {html.escape(resume)}" if resume else ""
        here = " ← you're here" if key == cur else ""
        lines.append(f"• <b>{html.escape(meta.get('label') or key)}</b>{tail}{here}")
    if cur == threads.MAIN:
        lines.append("")
        lines.append("You're in the <b>main</b> conversation.")
    lines.append("")
    lines.append("Typing continues the thread you're in. Reply to a message to "
                 "switch to another. /main leaves threads, /reset clears them.")
    await update.message.reply_text("\n".join(lines), parse_mode="HTML")


async def main_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    """Leave the current job thread and go back to the general conversation."""
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    threads.set_current(update.effective_chat.id, threads.MAIN)
    await update.message.reply_text(
        "↩️ Back in the main conversation. Reply to a job's message to return to it.")


async def on_message(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    chat_id = update.effective_chat.id
    pending_jid = _pending_skip_reason.pop(chat_id, None)
    if pending_jid:
        reason = (update.message.text or "").strip()
        jobs_store.set_decision(pending_jid, "skipped", reason or None)
        await update.message.reply_text("Got it — noted why you skipped. 👍")
        return
    if onboarding.status() == "not_started" and onboarding.is_fresh():
        await _launch_onboarding(update, ctx, user_message=update.message.text)
        return
    thread, opened = _route_thread(update)
    if opened:
        await _announce_thread(ctx.bot, chat_id, thread)
    prompt = update.message.text
    if thread != threads.MAIN:
        prompt += "\n\n" + JOB_MARKER_INSTRUCTION
    await _run_and_reply(update, ctx, prompt, thread=thread)


async def on_document(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    doc = update.message.document
    caption = (update.message.caption or "").strip()
    dest = config.UPLOADS_DIR / _safe_name(doc.file_name or f"{doc.file_unique_id}")
    try:
        tg_file = await ctx.bot.get_file(doc.file_id)
        await tg_file.download_to_drive(str(dest))
    except Exception as e:  # noqa: BLE001
        await update.message.reply_text(
            f"⚠️ Couldn't download that file ({_safe_error(e)}). Telegram caps bot downloads "
            "at ~20 MB — try a smaller file or paste the text.")
        return

    ext = dest.suffix.lower()
    note = f" The user added this note: {caption!r}." if caption else ""
    attach = None  # files attached for backends that need them (e.g. OpenCode --file)

    if ext == ".docx":
        extracted = _docx_to_text(dest)
        if not extracted:
            await update.message.reply_text(
                "I couldn't read that .docx. Please export it to PDF and resend, "
                "or paste the text.")
            return
        prompt = (f"The user uploaded their existing resume/CV ('{dest.name}'). "
                  f"Here is its full extracted text:{note}\n\n----\n{extracted}\n----\n\n"
                  "Extract their REAL experiences, skills, and profile from it and "
                  "save them into memory following your rules. Never invent anything "
                  "not present in the document. Then briefly summarize what you saved.")
    elif ext == ".pdf":
        extracted = await extract.extract_pdf(dest)
        if not extracted:
            await update.message.reply_text(
                "I couldn't read that PDF. If it's a scanned image, try a clearer "
                "copy, export it to a text-based PDF, or paste the text.")
            return
        prompt = (f"The user uploaded their existing resume/CV ('{dest.name}'). "
                  f"Here is its full extracted text:{note}\n\n----\n{extracted}\n----\n\n"
                  "Extract their REAL experiences, skills, and profile from it and "
                  "save them into memory following your rules. Never invent anything "
                  "not present in the document. Then briefly summarize what you saved.")
    elif ext in IMAGE_EXT:
        extracted = await extract.extract_image(dest)
        if not extracted:
            await update.message.reply_text(
                "I couldn't read that image. Try a clearer, well-lit copy, or "
                "paste the text.")
            return
        prompt = (f"The user uploaded an image of a document ('{dest.name}'), "
                  f"most likely their resume/CV.{note} Here is its transcribed "
                  f"text:\n\n----\n{extracted}\n----\n\n"
                  "Extract any REAL experiences, skills, or profile details and save "
                  "them into memory following your rules. Never invent anything not "
                  "present. Then briefly summarize what you saved and any gaps.")
    elif ext in NATIVE_READ_EXT:
        attach = [dest]
        prompt = (f"The user uploaded a document, saved at `uploads/{dest.name}` "
                  f"(most likely their existing resume/CV).{note} Read that file, then "
                  "extract their REAL experiences, skills, and profile and save them "
                  "into memory following your rules. Never invent anything not present "
                  "in the document. Then briefly summarize what you saved and any gaps.")
    else:
        await update.message.reply_text(
            f"I can read PDF, images, .txt/.md, and .docx. '{dest.name}' isn't one of "
            "those — please export it to PDF and resend.")
        return

    await _run_and_reply(update, ctx, prompt, files=attach,
                         thread=_route_thread(update, allow_new=False)[0])


async def on_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    if not _allowed(update):
        await update.message.reply_text("Sorry — this is a private bot.")
        return
    photo = update.message.photo[-1]  # largest size
    caption = (update.message.caption or "").strip()
    dest = config.UPLOADS_DIR / f"photo-{photo.file_unique_id}.jpg"
    try:
        tg_file = await ctx.bot.get_file(photo.file_id)
        await tg_file.download_to_drive(str(dest))
    except Exception as e:  # noqa: BLE001
        await update.message.reply_text(f"⚠️ Couldn't download that image ({_safe_error(e)}).")
        return

    note = f" The user added this note: {caption!r}." if caption else ""
    extracted = await extract.extract_image(dest)
    if not extracted:
        await update.message.reply_text(
            "I couldn't read that photo. Try a clearer, well-lit shot, or paste "
            "the text.")
        return
    prompt = (f"The user sent a photo of a document, saved as '{dest.name}' "
              f"(likely a resume, certificate, or document).{note} Here is its "
              f"transcribed text:\n\n----\n{extracted}\n----\n\n"
              "Extract any REAL experiences, skills, or profile details and save them "
              "into memory following your rules. Never invent anything not present. "
              "Then briefly summarize what you saved.")
    await _run_and_reply(update, ctx, prompt,
                         thread=_route_thread(update, allow_new=False)[0])


# Telegram's own transport wobbling (502s, read timeouts) during polling. These
# arrive with no update attached, mean nothing to the user, and were filling the
# log with full tracebacks — nine in three days.
_TRANSPORT_HICCUPS = (NetworkError, TimedOut)


async def on_error(update: object, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    """Last line of defence for anything a handler lets escape.

    Without this registered, python-telegram-bot logged the traceback to stdout
    and stopped there: the user sat looking at a chat where their message simply
    never got an answer. Same silent-failure shape as the dropped scan and the
    "(no response)" placeholder — the bot knew something had gone wrong and told
    nobody who cared.
    """
    e = ctx.error
    # Duck-typed rather than isinstance(update, Update): a non-Update simply
    # has no effective_chat, and the last thing this path should do is add an
    # assumption that can drop the very message it exists to deliver.
    chat = getattr(update, "effective_chat", None)

    if isinstance(e, _TRANSPORT_HICCUPS) and chat is None:
        # Polling hiccup with nobody waiting on it. One line, no traceback.
        log.warning("Telegram transport: %s", _safe_error(e))
        return

    log.error("Unhandled error while processing %s",
              type(update).__name__, exc_info=e)
    if chat is None:
        return
    try:
        await ctx.bot.send_message(
            chat_id=chat.id,
            text=f"⚠️ Something broke on my side: {_safe_error(e)}\n\n"
                 "Your message wasn't processed — try again, or /reset if it keeps happening.")
    except Exception:  # noqa: BLE001 — the notification failing must not recurse
        log.exception("Could not tell the user about the error above.")


def main() -> None:
    if not config.TELEGRAM_BOT_TOKEN:
        raise SystemExit("Set TELEGRAM_BOT_TOKEN in .env (get it from @BotFather).")
    if not (shutil.which(config.OPENCODE_BIN) or Path(config.OPENCODE_BIN).exists()):
        raise SystemExit(
            f"The OpenCode CLI ('{config.OPENCODE_BIN}') was not found. "
            "Install it and/or set OPENCODE_BIN. See docs/opencode-setup.md.")
    if not config.ALLOWED_USER_IDS:
        log.warning("ALLOWED_USER_IDS is empty — the bot is open to anyone who "
                    "finds it. Send /start to learn your ID, then lock it down.")

    async def _set_commands(application: Application) -> None:
        await application.bot.set_my_commands([
            ("scan", "Search for new job matches now"),
            ("status", "Show conversation size + reset button"),
            ("reset", "Start a fresh conversation (memory kept)"),
            ("help", "How to use me"),
            ("start", "Intro & your Telegram ID"),
            ("onboard", "Guided setup — re-run anytime"),
        ])

    app = (Application.builder()
           .token(config.TELEGRAM_BOT_TOKEN)
           .post_init(_set_commands)
           # Default is 1 — every update, however trivial, waited behind the
           # job thread ahead of it. Agent turns are still capped at
           # MAX_CONCURRENT_RUNS by the semaphore in agent.py; this only lets
           # button taps and /threads answer while a resume is being written.
           .concurrent_updates(config.MAX_CONCURRENT_UPDATES)
           .build())
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_cmd))
    app.add_handler(CommandHandler("reset", reset))
    app.add_handler(CommandHandler("onboard", onboard_cmd))
    app.add_handler(CommandHandler("status", status))
    app.add_handler(CommandHandler("threads", threads_cmd))
    app.add_handler(CommandHandler("main", main_cmd))
    app.add_handler(CallbackQueryHandler(on_callback))
    app.add_handler(MessageHandler(filters.Document.ALL, on_document))
    app.add_handler(MessageHandler(filters.PHOTO, on_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, on_message))
    app.add_handler(CommandHandler("scan", scan_cmd))
    app.add_error_handler(on_error)

    if config.JOB_DISCOVERY_ENABLED:
        if app.job_queue is None:
            log.warning("JobQueue unavailable — install python-telegram-bot[job-queue]. "
                        "Scheduled scans disabled; /scan still works.")
        else:
            app.job_queue.run_daily(
                _scheduled_scan,
                time=dt.time(hour=config.SCAN_HOUR, tzinfo=_SCAN_TZ),
                name="job-discovery-scan",
                # APScheduler defaults to a 1-second misfire window, so a busy
                # moment at the fire time silently skips the day's scan — that
                # happened on 2026-08-21, missed by 8.6s, with nothing but a
                # WARNING to show for it. A job scan does not care about being
                # late, so allow an hour. coalesce (already the default) keeps a
                # multi-day outage from firing a backlog of scans on restart.
                job_kwargs={"misfire_grace_time": config.SCAN_MISFIRE_GRACE},
            )
            log.info("Job discovery scheduled daily at %02d:00 %s on weekdays %s.",
                     config.SCAN_HOUR, config.SCAN_TZ, sorted(config.SCAN_WEEKDAYS))

    backend_desc = f"OpenCode (model: {config.OPENCODE_MODEL or 'default'})"
    log.info("Career Agent is running — backend: %s. Ctrl+C to stop.", backend_desc)
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
